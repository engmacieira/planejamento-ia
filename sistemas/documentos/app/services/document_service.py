import re
from docx import Document
from typing import List, Set, Dict

class DocumentService:
    def get_variables_from_file(self, file_path: str) -> List[str]:
        r"""
        Lê variáveis no formato ((tag)).
        Usa Negative Lookahead (?!\() para evitar capturar parênteses triplos (((...))).
        Ignora espaços em branco dentro da tag.
        """
        try:
            doc = Document(file_path)
            variables: Set[str] = set()
            
            # REGEX REFINADO:
            # \(\(      -> Busca 2 parênteses
            # (?!\()    -> GARANTE que não tem um terceiro parêntese logo em seguida
            # \s* -> Espaços opcionais
            # (.*?)     -> O conteúdo da variável
            # \s* -> Espaços opcionais
            # \)\)      -> Fecha 2 parênteses
            regex_pattern = r"\(\((?!\()\s*(.*?)\s*\)\)"

            for para in doc.paragraphs:
                matches = re.findall(regex_pattern, para.text)
                variables.update(matches)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            matches = re.findall(regex_pattern, para.text)
                            variables.update(matches)

            print(f"Variáveis encontradas: {list(variables)}")
            return list(variables)
            
        except Exception as e:
            print(f"Erro ao ler arquivo: {e}")
            return []
        
    def fill_document(self, template_path: str, data: Dict[str, str], output_path: str):
        """
        Abre o template, substitui as tags ((Chave)) pelos valores do dicionário 'data'
        e salva no output_path.
        """
        try:
            doc = Document(template_path)
            
            # Regex idêntico ao de leitura
            regex_pattern = r"\(\((?!\()\s*(.*?)\s*\)\)"

            def replace_text(text: str) -> str:
                # Função interna que faz a mágica da substituição
                matches = re.findall(regex_pattern, text)
                for match in matches:
                    # A chave encontrada (ex: "DFD")
                    key = match.strip()
                    
                    # Se tivermos o valor para essa chave, substituímos
                    if key in data:
                        value = str(data[key])
                        # Monta a tag original para busca (ex: "((DFD))" ou "(( DFD ))")
                        # O regex re.sub aqui é para garantir que substituímos a tag inteira
                        # incluindo os parênteses e espaços variáveis
                        pattern_replace = r"\(\((?!\()\s*" + re.escape(key) + r"\s*\)\)"
                        text = re.sub(pattern_replace, value, text)
                return text

            # 1. Substituir nos Parágrafos
            for para in doc.paragraphs:
                if "((" in para.text: # Otimização: só processa se tiver tag
                    para.text = replace_text(para.text)

            # 2. Substituir nas Tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if "((" in para.text:
                                para.text = replace_text(para.text)

            doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Erro ao preencher documento: {e}")
            raise ValueError(f"Falha ao gerar documento: {e}")