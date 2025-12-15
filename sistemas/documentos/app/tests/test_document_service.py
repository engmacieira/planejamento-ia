import pytest
import os
from docx import Document
from unittest.mock import patch
from app.services.document_service import DocumentService

@pytest.fixture
def temp_docx_path():
    """
    Cria um arquivo Word temporário simulando EXATAMENTE os cenários 
    que encontramos nos seus arquivos da prefeitura.
    """
    filename = "minuta_prefeitura_teste.docx"
    doc = Document()
    
    # 1. Caso Padrão: ((Tag))
    doc.add_paragraph("Processo administrativo nº ((DFD))")
    doc.add_paragraph("Objeto da licitação: ((Objeto))")
    
    # 2. Caso com espaços internos (Robustez): (( Tag ))
    # Isso garante que se alguém digitar com espaço sem querer, o sistema entende.
    doc.add_paragraph("Secretaria solicitante: (( Secretaria ))")
    
    # 3. Caso Crítico (O "Pesadelo" dos parênteses):
    # Cenário: Valor por extenso entre parênteses.
    # Texto no Word: "Valor: ((Valor)) (((ValorExtenso)))"
    # Interpretação: R$ 100,00 (cem reais)
    # Onde os parênteses externos são texto fixo, e os internos são a tag.
    doc.add_paragraph("Valor Global: ((Valor)) (((ValorExtenso)))") 
    
    # 4. Caso Tabela (Muitas minutas usam tabelas)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Data de Assinatura: ((Data1))"
    
    doc.save(filename)
    
    yield filename
    
    # Cleanup: Apaga o arquivo depois do teste para não deixar lixo
    if os.path.exists(filename):
        try:
            os.remove(filename)
        except PermissionError:
            pass # Às vezes o Windows segura o arquivo um pouco, sem problemas.

def test_extract_variables_custom_format(temp_docx_path):
    """
    Testa se o regex captura corretamente o padrão ((tag)),
    ignorando parênteses de texto ao redor e espaços internos.
    """
    service = DocumentService()
    
    # Act: Executa a leitura
    variables = service.get_variables_from_file(temp_docx_path)
    
    print(f"Variáveis encontradas pelo Regex: {variables}") # Debug útil no console
    
    # Assert (Validações)
    
    # 1. Deve encontrar as tags normais
    assert "DFD" in variables
    assert "Objeto" in variables
    
    # 2. Deve limpar espaços internos (( Secretaria )) -> "Secretaria"
    assert "Secretaria" in variables
    
    # 3. Deve lidar com o caso dos 3 parênteses (((ValorExtenso))) -> "ValorExtenso"
    # O regex deve ser esperto e pegar apenas o miolo, ignorando o parêntese de texto em volta.
    assert "ValorExtenso" in variables
    assert "Valor" in variables
    
    # 4. Prova Real: Não deve trazer os parênteses do texto junto
    # Se o regex for ganancioso demais, ele traria "(ValorExtenso)" -> Isso estaria errado.
    assert "(ValorExtenso)" not in variables 
    
    # 5. Deve encontrar tags dentro de tabelas
    assert "Data1" in variables
    
    # 6. Contagem final: DFD, Objeto, Secretaria, Valor, ValorExtenso, Data1 = 6 variáveis
    assert len(variables) == 6
    
def test_extract_variables_error_handling():
    """
    Testa se o service trata erros (ex: arquivo inexistente) 
    retornando lista vazia em vez de explodir.
    Cobre o bloco 'except Exception' (linhas 39-41).
    """
    service = DocumentService()
    
    # Act: Tenta ler um arquivo que não existe
    # Isso vai gerar um erro interno no python-docx, que cairá no except
    result = service.get_variables_from_file("arquivo_fantasma_que_nao_existe.docx")
    
    # Assert: Deve retornar lista vazia (comportamento seguro definido por nós)
    assert result == []
    assert isinstance(result, list)
    
def test_fill_document_success(temp_docx_path):
    """
    Testa se o service substitui as tags pelos valores fornecidos.
    """
    service = DocumentService()
    
    # 1. Dados para preencher (Baseado no seu CSV)
    dados = {
        "DFD": "038/2025",
        "Objeto": "Aquisição de Medicamentos",
        "Valor": "R$ 100,00",
        "ValorExtenso": "cem reais",
        "Data1": "20 de Outubro"
    }
    
    output_filename = "resultado_teste.docx"
    
    # 2. Act: Gerar o documento
    service.fill_document(temp_docx_path, dados, output_filename)
    
    # 3. Assert: Ler o documento gerado e verificar se os valores estão lá
    doc = Document(output_filename)
    full_text = " ".join([p.text for p in doc.paragraphs])
    
    # Debug
    print(f"Texto gerado: {full_text}")
    
    assert "038/2025" in full_text
    assert "Aquisição de Medicamentos" in full_text
    # Verifica aquele caso chato dos 3 parênteses:
    # Template: "Valor Global: ((Valor)) (((ValorExtenso)))"
    # Esperado: "Valor Global: R$ 100,00 (cem reais)"
    assert "R$ 100,00" in full_text
    assert "(cem reais)" in full_text
    assert "((Valor))" not in full_text # A tag deve ter sumido
    
    # Cleanup
    if os.path.exists(output_filename):
        os.remove(output_filename)
        
    from unittest.mock import patch

def test_fill_document_error_handling(temp_docx_path):
    """
    Força um erro durante o preenchimento para testar o bloco except.
    Cobre: raise ValueError(f"Falha ao gerar documento: {e}")
    """
    service = DocumentService()
    dados = {"DFD": "123"}
    
    # Mockamos a classe Document para lançar um erro quando for chamada
    with patch("app.services.document_service.Document", side_effect=Exception("Erro fatal do Word")):
        
        # Esperamos que o serviço capture o erro interno e lance um ValueError nosso
        with pytest.raises(ValueError) as excinfo:
            service.fill_document(temp_docx_path, dados, "saida.docx")
        
        # Verifica se a mensagem de erro foi preservada
        assert "Falha ao gerar documento" in str(excinfo.value)
        assert "Erro fatal do Word" in str(excinfo.value)